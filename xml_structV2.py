import string
import xml.etree.ElementTree as ET
# import docx.styles
# import docx.styles.style
import openpyxl as xl
import pandas as pd
import zipfile
import math
import json
import docx

from openpyxl.styles import PatternFill
from tkinter.filedialog import askopenfilenames
from docx.shared import Pt
# from docx.styles import style
# import placement_of_objects
import sql_data



"""
Нужно бы создать класс для помещений
"""
class Scan_Visio:
    """ Парсер vsdx документа """
    def __init__(self,files,choice,split_tables):
        """ Инициализация программы. Данный класс является корневым элементом.
        Здесь подгружаются все наименования объектов из файла lists_of_all_elements.json, встречаемые на схеме Visio
        Редактирование сканируемых элементов выполняется непосредственно в упомянутом прежде файле формата .json"""
        print('start script Scan_Visio')
        with open('lists_of_all_elements.json', encoding='utf-8') as elements:
            self.elements = json.load(elements,)
        try:
            self.split_tables = int(split_tables)
            if self.split_tables != 1:
                self.split_tables = 0
        except:
            self.split_tables = 0
        # инициализация стартовых данных
        self.choice = choice # была создана для определения необходимости создания определённой таблицы, в данный момент идея является устаревшей
        self.files = files
        self.index_start_name = [i for i, sym in enumerate(self.files[0]) if sym == '/'][-1]
        self.index_start_name += 1
        
        self.rows = list(string.ascii_uppercase) # столбцы для переноса данных в Excel
        
        # Инициализация запретных символов для корректного считывания элементов
        self.wrong_sym = list(string.ascii_lowercase)
        self.numbers = [str(num) for num in range(10)]
        self.numbers.append('.')

        self.master_nec_el = {}
        self.result = {} # хранит в себе мастер(число в строковом типе) - кол-во элементов  
        self.shapes_worklists = {} 
        
        self.db_init_()
        self.room_characteristics()  
        self.init_ITSO()
        self.init_POU_KEO()
        self.init_other_elements()
        


    def db_init_(self):
        """ Инициализация базы данных, для регистрации добавленных элементов программно,
          а так же собранной информации о уже присутствующих элементах"""
        self.sql = sql_data.Base_data(self.files[0][:self.index_start_name])
        self.tables = self.sql.get_info_tables()


    def init_ITSO(self):
        """ Подгрузка наименований типов ИТСО, для корректного считывания и ранжирования"""
        # Сделать подгрузку данных не привязанной к определённым типам. Тобишь универсалом. СМ. все иниты элементов
        self.video_elements = self.elements['video_elements']
        self.fire_elements = self.elements['fire_elements']
        self.security_elements = self.elements['security_elements']
        self.necessary_elements = *self.video_elements, *self.fire_elements, *self.security_elements
        

    def init_POU_KEO(self):
        """ Подгрузка наименований типов ПОУ и КЭО, для корректного считывания"""
        # Сделать подгрузку данных не привязанной к определённым типам. Тобишь универсалом. СМ. все иниты элементов

        self.table_KEO = [] # тут хранятся все элементы КЭО
        self.table_POU = [] # ТУТ хранятся все эементы ПОУ
        self.table_POU_KEO = {file: {'КЭО':[],'ПОУ':[]} for file in self.files} # тут хранятся "файл - ((айди, значение элемента),(айди, значение элемента))"
        self.masters_pou_KEO = {'КЭО':0,
                        'ПОУ':0}
        self.place_on_floors = self.elements['place_on_floors']
        for file in self.files:
            with zipfile.ZipFile(file, 'r') as zf:
            # Извлечение содержимого XML файла
                with zf.open('visio/pages/page1.xml', 'r') as f:
                # Разбор XML документа
                    self.tree = ET.parse(f)
                    self.root = self.tree.getroot()
                    for child in self.root[0]:
                        elem = child.attrib
                        try:
                            #  получаем опасные участки на схеме
                            if 'ПОУ' in elem['NameU']:
                                self.masters_pou_KEO['ПОУ'] = elem['Master']
                            elif 'КЭО' in elem['NameU']:
                                self.masters_pou_KEO['КЭО'] = elem['Master']
                        except KeyError:
                            pass


    def init_other_elements(self):
        ''' Подгрузка остальных наименований(3-я таблица). '''
        self.other_elements = self.elements['other_elements']
        self.other_table_elements = []

    def checking_compatibility(self):
        """ Проверка присутствия всех типов наименований на схеме
        Данный метод сверяет элементы из подгруженных наименований с джейсона и, в случае отсутствия элемента сообщает этот элемент пользователю.
        Данный метод не является методом постоянного использования. Юзается исключительно для проверки совместимости файла Visio и программы не чаще чем раз в месяц"""
        # Вероятно следует создать новый файл с отдельным назначением, либо создать тумблер на проверку файла, либо совместить со сканированием всех мастеров СМ file_masters, а так же цикл в init_POU_KEO
        self.get_masters()
        all_elements = *self.other_elements, *self.necessary_elements, *self.place_on_floors
        key_masters = self.masters.keys()
        for key in all_elements:
            if key not in key_masters:
                print(key, "Не обнаружен, обновите версию чертежа")

    def file_masters(self,zf):
        with zf.open('visio/masters/masters.xml', 'r') as f:
                    tree = ET.parse(f)
                    root = tree.getroot()
                    for element in root:
                        try:
                            self.master_keys[element.attrib['Name']] = element.attrib['ID']
                        except KeyError:
                            pass

    def file_shape(self,zf): 
        """ Получение размеров чертежа по ширине и высоте соответственно.
        Сие данные юзаются в файле visualization, для дальнейшего отображения элементов"""
        with zf.open('visio/pages/pages.xml', 'r') as f:
            self.tree = ET.parse(f)
            self.root = self.tree.getroot()[0][0]
            w = float(self.root[0].attrib['V'])
            h = float(self.root[1].attrib['V'])

            self.shapes_worklists[self.file[self.index_start_name:-5]] = w,h

    

    def check_cheme(self):
        """ Корневой метод.
        Парсит файлы по списку и сверяется с уже известными признаками
        Парс данные делятся на типы: ИТСО, ПОУ И КЭО, Иные элементы - 
        Стены двери, лестницы и плиты(указание глухих помещений) парсятся для дальнейшего отображения на личном приложении
        На перспективу: Можно пихнуть в отельный метод все 3 разные проверки стен, дверей итд
        
        """
        # self.checking_compatibility() # - для проверки совместимости. НЕ ТРОГАТЬ
        for key in self.necessary_elements: # Подготовка заполнительного поля на каждый файл для мастеров и кол-ва подсчёта
            self.result[key] = list(map(int,list('0' * (len(self.files)+1))))
            self.master_nec_el[key] = list(map(int,list('0' * len(self.files))))

        for index_file, self.file in enumerate(self.files): # парс файлов
            self.count_elements = {}
            self.master_keys = {}
            
            # Открытие ZIP архива внутри VSDX файла
            with zipfile.ZipFile(self.file, 'r') as zf:
                self.file_masters(zf)
                self.file_shape(zf)
                # начало парсинга xml структуры файла для поиска и подсчёта элементов
                with zf.open('visio/pages/page1.xml', 'r') as f:
                    self.tree = ET.parse(f)
                    self.root = self.tree.getroot()
                    
                    # С "корневого" элемента начинаем парсить все дочерние элементы
                    for self.index_elem, self.child in enumerate(self.root[0]): 
                        self.elem = self.child.attrib
                        try:
                            if self.choice['ПОУ и КЭО']:
                                try:
                                    self.check_POU_KEO()
                                except KeyError:
                                    pass
                                
                            if self.elem['Master'] not in self.master_keys.values():
                                if self.elem['NameU'].isalpha(): # исключение вероятности появления чисел в наименовании
                                    self.master_keys[self.elem['NameU']] = self.elem['Master']
                                    if 'Wall' in self.elem['NameU']: 
                                        self.transfer_wall()
                                    elif 'Curved wall' in self.elem['NameU']:
                                        self.transfer_wall()
                                    elif 'Straight staircase' in self.elem['NameU']:
                                        self.transfer_stair()
                                    elif 'Door' in self.elem['NameU'] or 'Double' in self.elem['NameU'] or 'Opening' in self.elem['NameU']:
                                        self.transfer_door()
                                    elif 'Slab' in self.elem['NameU']:
                                        self.transfer_pilaster()
                                    
                                    self.count_elements[self.elem['Master']] = 0
                                else: # В случае, если наименование содержит неверные символы - мы их удаляем и проверяем совпадение с искомым
                                    # Данный способ можно так же вынести в отдельный метод для удобночитаемости
                                    wrong_word = list(self.elem['NameU'])
                                    i = 0
                                    while 1:
                                        try:
                                            sym = wrong_word[i]
                                        except:
                                            break
                                        if sym in self.numbers:
                                            wrong_word.remove(sym)
                                        else:
                                            i += 1
                                    good_word = ''.join(wrong_word)
                                    self.master_keys[good_word] = self.elem['Master']
                                    self.count_elements[self.elem['Master']] = 0
                                    # Можно выкинуть в отдельный метод
                                    if 'Wall' in self.elem['NameU']:
                                        self.transfer_wall()
                                    elif 'Curved wall' in self.elem['NameU']:
                                        self.transfer_wall()
                                    elif 'Straight staircase' in self.elem['NameU']:
                                        self.transfer_stair()
                                    elif 'Door' in self.elem['NameU']  or 'Double' in self.elem['NameU'] or 'Opening' in self.elem['NameU']:
                                        self.transfer_door()
                                    elif 'Slab' in self.elem['NameU']:
                                        self.transfer_pilaster()

                            else:
                                if self.elem['Master'] in self.master_keys.values():
                                    """  с целью предупреждения редкой ошибки проверяем дважды"""
                                    key_master = self.get_key(self.elem['Master'])

                                    if 'Стена' in key_master or 'Кривая стена' in key_master:
                                        self.transfer_wall()
                                    elif 'Прямая лестница' in key_master:
                                        self.transfer_stair()
                                    elif 'Дверь' in key_master or 'Двойная' in key_master or 'Проём' in key_master:
                                        self.transfer_door()
                                    elif 'Плита' in key_master: # Плита - предмет которым заполняем помещения которых нет
                                        self.transfer_pilaster()
                                    else:
                                        try:
                                            self.count_elements[self.elem['Master']] += 1
                                        except KeyError:
                                            try:
                                                self.count_elements[str(self.elem['Master'])] = int(0)
                                            except:
                                                pass
                        except KeyError:
                            pass

            self.upload_rooms()
            for val in self.master_keys:
                if val in self.result.keys():
                    try:
                        self.result[val][index_file] = self.count_elements[self.master_keys[val]]
                    except:
                        pass
                    if val in self.master_nec_el.keys():
                        self.master_nec_el[val][index_file] = self.master_keys[val]
                try:
                    if val in self.other_elements and self.count_elements[self.master_keys[val]] > 0 and val not in self.other_table_elements:
                        self.other_table_elements.append(val)
                except Exception as e: 
                        print(f'Исключение !!! {val}', e)
                    
    
        """ Добавлено определение иных элементов но нет вывода в next_etap"""
        if self.other_table_elements != []:
            self.sql.cashe('other_elements',self.other_table_elements)

        result = self.result.copy()
        for key, row, in self.result.items():
            if result[key] == [0] * len(row):
                del result[key]

        if self.choice['ИТСО'] or self.choice['Подсчёт']:
            excel = Excel(self.files)
            excel.create_table(self.result)

        self.sql.cashe('itso', [result])
        if self.choice['ПОУ и КЭО']:
            self.sort_dangerous()
            
        pos_tables = {file[0][self.index_start_name:-5] : [None,None,None,None] for file in self.files}
        # self.sql.check_db()
        # self.sql.update()
        # self.sql.close()
        return self.files_list_of_rooms_charact,self.shapes_worklists, result.keys(), self.table_POU_KEO , self.other_table_elements, self.split_tables,self.sql, self.files, pos_tables

    def sort_dangerous(self):
        """ Сортировка элементов КЭО и отдельно ПОУ, для нумераци в порядке алфавита"""
        keo, pou = [], []
        for file in self.table_POU_KEO.keys():
            for type in self.table_POU_KEO[file].keys():
                for name in self.table_POU_KEO[file][type]:
                    if 'КЭО' in name[2]:
                        keo.append(name[2])
                    elif 'ПОУ' in name[2]:
                        pou.append(name[2])
        keo.sort()
        for i,element in enumerate(keo):
            keo[i] = element, i+1
        
        pou.sort()
        for i,element in enumerate(pou):
            pou[i] = element, i+1
        
        self.table_POU_KEO['all_table'] = keo + pou
        excel = Excel(self.files)
        excel.add_Dangerous(self.table_POU_KEO['all_table'])
        self.sql.cashe('dangerous',self.table_POU_KEO)
    
    def get_masters(self):
        """ Считывает мастера для проверки совместимости файла с версией ПО"""
        """Аналогичный метод с данным именем в placement_of_objects(файл для размещения фигур)""" 
        with zipfile.ZipFile(self.files, 'a') as visio:
            self.masters = {}
            with visio.open('visio/masters/masters.xml', 'r') as masters:
                tree = ET.parse(masters)
                masters_root = tree.getroot()
                for name_element in masters_root:
                    try:
                        name_element = name_element.attrib
                        self.masters[name_element['Name']] = name_element['ID']
                    except KeyError as e:
                        print('Мастера не были получены. : get_masters', e)


    def get_key(self, search_value):
        """ Получение имени элемента исходя из его мастера"""
        for key, value in self.master_keys.items():
            if value == search_value:
                return key
        return ''
            

    def check_POU_KEO(self):
        """ Проверка является ли элемент элементом поу или кэо
        Тут же фиксируется в каком типе тега заключен текст для его дальнейшего редактирования"""
        if 'Text' in self.child[-1].tag:
            accept_POU = False
            accept_KEO = False
            for index_main_text, line in enumerate(self.child[-1]):
                line_text = line.text if line.text != None else ''
                line_tail = line.tail if line.tail != None else ''
                
                if 'ПОУ' in line_text.upper() and 'КЭО' not in line_text.upper() or 'ПОУ' in line_tail.upper() and 'КЭО' not in line_tail.upper():
                    if len(line_text) > 6:
                        self.table_POU_KEO[self.file]['ПОУ'].append([self.elem['ID'],self.index_elem, line_text, 'text',index_main_text])
                        break
                    elif len(line_tail) > 6:
                        self.table_POU_KEO[self.file]['ПОУ'].append([self.elem['ID'],self.index_elem, line_tail, 'tail',index_main_text])
                        break
                    else:
                        accept_POU = True
                        continue

                if accept_POU and len(line_text) > 2:
                    self.table_POU_KEO[self.file]['ПОУ'].append([self.elem['ID'],self.index_elem, line_text, 'text',index_main_text])
                    break
                elif accept_POU and len(line_tail) > 2:
                    self.table_POU_KEO[self.file]['ПОУ'].append([self.elem['ID'],self.index_elem, line_tail, 'tail',index_main_text])
                    break
                
                if 'КЭО' in line_text.upper() and 'ПОУ' not in line_text.upper() or 'КЭО' in line_tail.upper() and 'ПОУ' not in line_tail.upper():
                    if len(line_text) > 6:
                        self.table_POU_KEO[self.file]['КЭО'].append([self.elem['ID'],self.index_elem, line_text, 'text',index_main_text])
                    elif len(line_tail) > 6:
                        self.table_POU_KEO[self.file]['КЭО'].append([self.elem['ID'],self.index_elem, line_tail, 'tail',index_main_text])
                    else:
                        continue

                if accept_KEO and len(line_text) > 2:
                    self.table_POU_KEO[self.file]['КЭО'].append([self.elem['ID'],self.index_elem, line_text, 'text',index_main_text])
                    break
                elif accept_KEO and len(line_tail) > 2:
                    self.table_POU_KEO[self.file]['КЭО'].append([self.elem['ID'],self.index_elem, line_tail, 'tail',index_main_text])
                    break
            else: # условие для "обычного" кэо
                if self.child[-1].text != None:
                    if 'КЭО' in self.child[-1].text:
                        self.table_POU_KEO[self.file]['КЭО'].append([self.elem['ID'],self.index_elem, self.child[-1].text, 'text',None])
                    elif 'ПОУ' in self.child[-1].text:
                        self.table_POU_KEO[self.file]['ПОУ'].append([self.elem['ID'], self.index_elem, self.child[-1].text, 'text',None])

    

    def room_characteristics(self):
        """ Метод инициализации для обработки информации о стенах, дверях, лестницах и плитах
        Все переменные из данного метода используются в файле visualization для нумерации и экспликации помещений"""
        self.files_list_of_rooms_charact = {file[self.index_start_name:-5]:{} for file in self.files}
        self.list_of_walls = []   
        self.list_of_doors = []   
        self.list_of_stairs = []  
        self.list_of_pilaster = [] 


    def transfer_pilaster(self):
        """ Сканирование параметров плиты, используемой для обозначения глухих помещений"""
        """
        Позже перерасмотреть данный метод, нам нужен только центр данной фигуры для удаления формы помещения
        """
        x, y, center_x, center_y = None, None, None, None
        width, height = None, None
        for elem in self.child:
            if elem.attrib['N'] == "PinX":
                x = float(elem.attrib['V'])
            if elem.attrib['N'] == "PinY":
                y = float(elem.attrib['V'])
            if elem.attrib['N'] == 'Width':
                width = float(elem.attrib['V']) // 2
                x -= width
            if elem.attrib['N'] == 'Height':
                height = float(elem.attrib['V']) // 2
                y -= height
            if elem.attrib['N'] == 'Control':
                for child_elem in elem:
                    if child_elem.attrib['N'] == 'Row_2':
                        center_x = float(child_elem[0].attrib['V'])
                        center_y = float(child_elem[1].attrib['V'])
                        break
        if center_x == None: center_x = 50.
        if center_y == None: center_y = 50.
        x += center_x
        y += center_y
        self.list_of_pilaster.append([x,y])


    def transfer_wall(self): # получаем данные о схеме здания(корды, ширина каждой стены)
        """ Сканирование параметров стены """
        x0, y0, x1, y1 = None, None, None, None
        for i,elem in enumerate(self.child):
            if elem.attrib['N'] == "BeginX" and x0 == None:
                x0 = float(self.child[i].attrib['V'])
                y0 = float(self.child[i+1].attrib['V'])
            if elem.attrib['N'] == 'EndX' and x1 == None:
                x1 = float(self.child[i].attrib['V'])
                y1 = float(self.child[i+1].attrib['V'])
            if elem.attrib['N'] == "User":
                self.width_wall = 3.93700787402
                break
            if elem.attrib['N'] == "Property":
                try:
                    self.width_wall = float(self.child[i][0][0].attrib['V'])
                except:
                    self.width_wall = 3.93700787402
                break
        if self.width_wall > 15.: self.width_wall = 3.93700787402
        self.list_of_walls.append([x0,y0,x1,y1,self.width_wall])


    def transfer_door(self):
        """ Сканирование параметров двери"""
        if self.child[3].attrib['N'] == 'Height':
            height = 3.937007874015748
        else:
            height = float(self.child[3].attrib['V']) 
        x,y, h_wall = float(self.child[0].attrib['V']), float(self.child[1].attrib['V']), height
        self.list_of_doors.append([x,y,h_wall])        


    def transfer_stair(self):
        """ Сканирование параметров лестницы"""
        x,y,angle = float(self.child[0].attrib['V']), float(self.child[1].attrib['V']), 0.
        width = None
        for param in self.child:
            if param.attrib['N'] == 'Width':
                width = float(param.attrib['V'])
            try:
                if param.attrib['N'] == 'Angle':
                    angle = float(param.attrib['V'])
            except KeyError:
                pass
            
            if param.attrib['N'] == 'Property':
                if param[-1][0].attrib['V'] == '1':
                    angle += 0
                    if width is None:
                        x = 50 * math.cos(angle) + x
                        y = 50 * math.sin(angle) + y
                    else:
                        x = width//2 * math.cos(angle) + x
                        y = width//2 * math.sin(angle) + y
                break

        self.list_of_stairs.append([x, y, angle])
        

    def upload_rooms(self): 
        """ Закрепляем сканированные данные за определённым чертежём"""
        self.files_list_of_rooms_charact[self.file[self.index_start_name:-5]]['walls'] = self.list_of_walls
        self.files_list_of_rooms_charact[self.file[self.index_start_name:-5]]['doors'] = self.list_of_doors
        self.files_list_of_rooms_charact[self.file[self.index_start_name:-5]]['stairs'] = self.list_of_stairs
        self.files_list_of_rooms_charact[self.file[self.index_start_name:-5]]['slab'] = self.list_of_pilaster
        self.list_of_walls = []
        self.list_of_doors = []
        self.list_of_stairs = []
        self.list_of_pilaster = []

class Excel:
    def __init__(self,files):
        self.columns = list(string.ascii_uppercase)
        self.files = files
        self.elements = {'ИТСО': {}, 'ПОУ и КЭО': {'ПОУ': [], 'КЭО': []}}

    def add_Dangerous(self, data):
        xl_path = self.files[0][:get_dir(self.files[0])]+'Подсчёт.xlsx'
        wb = xl.load_workbook(xl_path)
        ws = wb.active
        row_index = 2
        
        while ws['A'+str(row_index)].value != None:
            row_index += 1
        
        row_index += 1
        ws['A'+str(row_index)].value = "Таблица ПОУ и КЭО"
        row_index += 1
        first_type = data[0][0][:3]
        ws['A'+str(row_index)].value = first_type
        
        for element in data:
            element_type, name = element[0][:3], element[0][4:]
            row_index += 1
        
            if element_type != first_type:
                ws['A'+str(row_index)].value = element_type
                row_index += 1
                first_type = element_type
                
            ws['A'+str(row_index)].value = name
        
        ws['A'+str(row_index+1)].value = '--------'
        
        wb.save(xl_path)
        
    def get_all_info(self, file):
        ''' Получение инфо для трансфера в word'''
        wb = xl.load_workbook(file)
        ws = wb.active
        cell = ''
        sym_index = 0
        
        while cell != 'Итог':
            sym_index += 1
            cell = ws[self.columns[sym_index]+'1'].value
        row = 2
        
        column_names_object = self.columns[sym_index+1]
        while ws['A'+str(row)].value != None:
            count = 0
            for column_count in self.columns[1:sym_index]:
                f = ws[column_count+str(row)].value
                try:
                    count += int(ws[column_count+str(row)].value)
                except: 
                    pass
                
            names = []
            for name_object in ws[column_names_object+str(row)].value.split(','):
                if name_object[0] == ' ':
                    name_object = name_object[1::]
                if name_object[-1] == ' ':
                    name_object = name_object[:-1:]
                if name_object != '-':
                    names.append(f' «{name_object}»') 
                else:
                    names.append('')
                    
            line = ''.join(names) + f" – {count} шт.,"
            self.elements['ИТСО'][ws['A'+str(row)].value] = line 
            row += 1
            
        row += 1
        if ws['A'+str(row)].value == 'Таблица ПОУ и КЭО':
            row += 1
            first_type = ws['A'+str(row)].value
            row += 1
            while ws['A'+str(row)].value != '--------':
                if ws['A'+str(row)].value in self.elements['ПОУ и КЭО'].keys():
                    first_type = ws['A'+str(row)].value
                    row += 1
                self.elements['ПОУ и КЭО'][first_type].append(ws['A'+str(row)].value)
                row += 1
        return self.elements
    
    def init_types(self):
        with open('lists_of_all_elements.json', encoding='utf-8') as elements:
            self.elementos = json.load(elements,)
        self.video_elements = self.elementos['video_elements']
        self.fire_elements = self.elementos['fire_elements']
        self.security_elements = self.elementos['security_elements']

    def create_table(self, elements):
        """ Создаётся таблица формата xlsx с указанием всех элементов ИТСО"""
        self.init_types()
        self.result = elements
        self.index_start_name = get_dir(self.files[0])
        index_rows = [name[self.index_start_name:-5] for name in self.files]
        index_rows.append('Итог')
        file = f'{self.files[0][:self.index_start_name]}Подсчёт.xlsx'
        self.data = {}
        list_deleted = []
        
        for key in self.result.keys():
            if self.result[key] == list(map(int,list('0' * (len(self.files)+1)))):
                list_deleted.append(key)

        for key in list_deleted:
            self.result.pop(key)
            
        for i,val in enumerate(self.result.keys()):
            try:
                self.result[val][-1] = f'=SUM({str(self.columns[1])}{str(i+2)}:{str(self.columns[len(self.files)])}{str(i+2)})'
            except IndexError:
                self.result[val][-1] = f'=SUM({str(self.columns[1])}{str(i+2)}:{str(self.columns[0])}{str(self.columns[len(self.files)-24])}{str(i+2)})'
        
        v,f,s = 0,0,0
        
        for key in self.result:
            if key in self.video_elements:
                v += 1
                self.data[key] = self.result[key]
        for key in self.result:
            if key in self.fire_elements:
                f += 1
                self.data[key] = self.result[key]
        for key in self.result:
            if key in self.security_elements:
                s += 1
                self.data[key] = self.result[key]

        df = pd.DataFrame(self.data,index=index_rows).T
        df.to_excel(file)

        wb = xl.load_workbook(file)
        ws = wb.active
        ws.column_dimensions['A'].width = 60
        
        if v != 0:
            for num in range(2, v+2):
                ws['A'+str(num)].fill = PatternFill('solid', fgColor="0070C0")
        
        if f != 0:
            
            for num in range(v+2,v+f+2):
                ws['A'+str(num)].fill = PatternFill('solid', fgColor="FF0000")
        
        if s != 0:
            for num in range(v+f+2, v+f+s+2):
                ws['A'+str(num)].fill = PatternFill('solid', fgColor="00B050")
        
        wb.save(file)


class Word:
    """ Работа с форматом .docx
    В настоящий момент юзается для переноса данных из Excel в Word"""
    def __init__(self, file, document = None):
        self.doc = docx.Document(document)
        self.load_dict()   
        self.dir = file[:get_dir(file)]
        
    def open_file(self):
        self.doc = docx.Document(f'{self.dir}info.docx')        
         
    def load_dict(self):
        ''' Загрузка словаря для корректного разделения элементов на типы '''
        with open('lists_objects_for_word.json', encoding='utf-8') as f:
            self.words = json.load(f,)
    
    def write_data_Dangerous(self, data):
        ''' Мост передачи ПОУ и КЭО '''
        for type_element in data['ПОУ и КЭО']:
            self.doc.add_heading(f'Элементы {type_element}')
            for element in data['ПОУ и КЭО'][type_element]:
                paragraph = self.doc.add_paragraph(element)
                style_paragraph = paragraph.runs[0]
                style_paragraph.font.name = 'Times New Roman'
                style_paragraph.font.size = Pt(8)
                
    def write_data_ITSO(self, data):
        ''' Мост передачи ИТСО '''
        for type_chapter in self.words.keys():
            heading = self.doc.add_heading(type_chapter)
            elements = []
            for type_element in self.words[type_chapter]:
                if type_element in data['ИТСО'].keys():
                    if 'пож' in type_element:
                        k = type_element.split(' ')
                        try:
                            k[k.index('пож')] = 'пожарный'
                            element = ' '.join(k)
                        except:
                            element = type_element
                    else:
                        element = type_element

                    if type_element[1].islower():
                        elements.append(f"{element[0].lower() + element[1::]} {data['ИТСО'][type_element]}")
                    else:
                        elements.append(f"{type_element} {data['ИТСО'][type_element]}")

                    f = type_element    
                    del data['ИТСО'][type_element]

            paragraph = self.doc.add_paragraph(' '.join(elements))
        
        paragraph.style.font.name = 'Times New Roman'
        paragraph.style.font.size = Pt(12)
        self.doc.add_heading('Элементы не вошедшие ни в 1 из категорий: ')
        elements = []
        line = ''

        for type, name in data['ИТСО'].items():
            line.join(type, )
            if type[1].islower():
                elements.append(f" {type[0].lower() + type[1::]} {name}")
            else:
                elements.append(f" {type} - {name}")

        self.doc.add_paragraph(''.join(elements))


def get_dir(path):
    ''' Получение текущей директории '''
    end_file_dir = 0
    for index, sym in enumerate(path):
        if sym == '/':
            end_file_dir = index
            
    return end_file_dir + 1


def start(choice):
    """Старт парсинга файла с возвратом данных"""
    if choice['Transfer']:
        files = askopenfilenames(filetypes=[('','xlsx')])
        excel = Excel(files)
        word = Word(files[0])
        excel_info = excel.get_all_info(files[0])
        word.write_data_ITSO(excel_info)
        word.write_data_Dangerous(excel_info)
        word.doc.save(f'{word.dir}info.docx')
    else:
        files = askopenfilenames(filetypes=[('', 'vsdx')])
        split = 1
        if len(files) != 0:
            print(files)
            Visio = Scan_Visio(files,choice,split)
            return Visio.check_cheme()
